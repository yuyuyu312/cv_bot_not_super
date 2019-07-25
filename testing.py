import telebot;
from telebot import types
from docx import Document
from docx.shared import Inches

document = Document()

bot = telebot.TeleBot('858195729:AAH43M75DLC0G8BkN3nfFAfs3ZHhWepQzIY');

full_name = '';
adress = '';
age = 0;


@bot.message_handler(content_types=['text'])
def start(message):
    if message.text.lower() == 'hi':
        user_name_data = str(message.from_user)
        bot.send_message(message.from_user.id, "Hi! Welcome to CV Bot.\nPlease, enter Your full name : ");
        bot.register_next_step_handler(message, get_full_name);

    else:
        bot.send_message(message.from_user.id, 'write /hi');


def get_full_name(message):
    global first_name;
    full_name_data = full_name = message.text;
    bot.send_message(message.from_user.id, "Please, enter Your adress: ");
    bot.register_next_step_handler(message, get_adress);
    document.add_heading(full_name_data)


def get_adress(message):
    global last_name;
    adress_data = adress = message.text;
    bot.send_message(message.from_user.id, "Please, enter Your age: ");
    bot.register_next_step_handler(message, get_age)
    document.add_heading(adress_data)


def get_age(message):
    global age;
    # while age == 0:
    try:
        age = int(message.text)
        bot.send_message(message.from_user.id, "Please, enter Your email: ");
        bot.register_next_step_handler(message, get_email)
        document.heading('age: ' + age, 3)

    except Exception:
        bot.send_message(message.from_user.id, 'Please, use numbers: ');


def get_email(message):
    global email;
    email_data = email = str(message.text);
    bot.send_message(message.from_user.id, "Please, enter Your phone number: ")
    bot.register_next_step_handler(message, get_phone_number)
    document.add_heading('Email: ' + email_data,  level=1)


def get_phone_number(message):
    global phone_number;
    phone_number_data = phone_number = int(message.text);
    bot.send_message(message.from_user.id, "Please, enter Your school's name: ")
    bot.register_next_step_handler(message, get_school_name)
    document.add_paragraph('phone_number: ' + str(phone_number_data))


def get_school_name(message):
    global school_name;
    school_name_data = school_name = str(message.text);
    bot.send_message(message.from_user.id, "Date enter school: (example: 31.12.2020)")
    bot.register_next_step_handler(message, get_enter_school_date)
    document.add_paragraph('scool name: ' + school_name_data)

def get_enter_school_date(message):
    global enter_school_date;
    enter_school_date_data = enter_school_date = str(message.text);
    bot.send_message(message.from_user.id, "Graduation date: (example: 31.12.2020)")
    bot.register_next_step_handler(message, get_graduation_year)
    document.add_paragraph('date of enter to school: ' + enter_school_date_data)


def get_graduation_year(message):
    global graduation_year;
    graduation_year_data = graduation_year = str(message.text);
    bot.send_message(message.from_user.id, "Please, enter wished job title: ")
    bot.register_next_step_handler(message, get_job_title)
    document.add_paragraph('yerf of graduate: ' + graduation_year_data)



def get_job_title(message):
    global job_title;
    job_title_data = job_title = str(message.text);
    keyboard = types.InlineKeyboardMarkup();
    key_yes = types.InlineKeyboardButton(text='Yes', callback_data='yes');
    keyboard.add(key_yes);
    key_no = types.InlineKeyboardButton(text='No', callback_data='no');
    keyboard.add(key_no);
    question = 'You are ' + full_name + ', ' + str(age) + ' years old, Wished job title is: '\
               + job_title + '. ' + 'Want to save it?'
    bot.send_message(message.from_user.id, text=question, reply_markup=keyboard)
    with open('cv.doc', 'a') as f:
        f.write(str(job_title_data) + '\n')


@bot.callback_query_handler(func=lambda call: True)
def callback_worker(call):
    if call.data == 'yes':
        bot.send_message(call.message.chat.id, 'Remember it: )')

    elif call.data == 'no':
        bot.send_message(call.message.chat.id, 'Delete it: )')


bot.polling(none_stop=True, interval=0)
document.save('demo.docx')
